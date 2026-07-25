from __future__ import annotations
from pathlib import Path
import json, html
import pandas as pd
import numpy as np
from ..config import guardar_configuracion

def _texto_celda(v):
    if isinstance(v, (list, tuple, dict, np.ndarray)):
        return json.dumps(v, ensure_ascii=False, default=str)
    try:
        if pd.isna(v): return ''
    except Exception:
        pass
    return str(v)

def exportar_excel(tablas:dict[str,pd.DataFrame],ruta):
    ruta=Path(ruta); ruta.parent.mkdir(parents=True,exist_ok=True)
    with pd.ExcelWriter(ruta,engine='openpyxl') as w:
        for nombre,df in tablas.items(): df.to_excel(w,sheet_name=nombre[:31],index=False)
    return ruta

def exportar_csv(tablas:dict[str,pd.DataFrame],carpeta):
    carpeta=Path(carpeta); carpeta.mkdir(parents=True,exist_ok=True); rutas=[]
    for n,df in tablas.items():
        p=carpeta/(n.replace(' ','_')+'.csv'); df.to_csv(p,index=False); rutas.append(p)
    return rutas

def generar_html(titulo,config,tablas,figuras,ruta,notas=None):
    ruta=Path(ruta); ruta.parent.mkdir(parents=True,exist_ok=True)
    partes=[f'<!doctype html><meta charset="utf-8"><title>{html.escape(titulo)}</title>', '<style>body{font-family:Arial;max-width:1200px;margin:auto;padding:24px}table{border-collapse:collapse;width:100%;margin-bottom:28px}th,td{border:1px solid #ccc;padding:6px}th{background:#eee}img{max-width:100%}pre{white-space:pre-wrap}</style>',f'<h1>{html.escape(titulo)}</h1>']
    if notas: partes.append(f'<p>{html.escape(notas)}</p>')
    partes+=['<h2>Configuración</h2>',f'<pre>{html.escape(json.dumps(config,ensure_ascii=False,indent=2))}</pre>']
    for n,df in tablas.items(): partes += [f'<h2>{html.escape(n)}</h2>',df.to_html(index=False,float_format=lambda x:f'{x:.6g}')]
    if figuras:
        partes.append('<h2>Gráficas</h2>')
        for p in figuras:
            rel=Path(p).resolve().relative_to(ruta.parent.resolve()) if Path(p).resolve().is_relative_to(ruta.parent.resolve()) else Path(p)
            partes.append(f'<figure><img src="{html.escape(str(rel))}"><figcaption>{html.escape(Path(p).stem)}</figcaption></figure>')
    ruta.write_text('\n'.join(partes),encoding='utf-8'); return ruta

def generar_docx(titulo,config,tablas,figuras,ruta):
    from docx import Document
    from docx.shared import Inches
    doc=Document(); doc.add_heading(titulo,0); doc.add_heading('Configuración',level=1); doc.add_paragraph(json.dumps(config,ensure_ascii=False,indent=2))
    for nombre,df in tablas.items():
        doc.add_heading(nombre,level=1); tabla=doc.add_table(rows=1,cols=len(df.columns)); tabla.style='Table Grid'
        for i,c in enumerate(df.columns): tabla.rows[0].cells[i].text=str(c)
        for _,row in df.iterrows():
            cells=tabla.add_row().cells
            for i,v in enumerate(row): cells[i].text=_texto_celda(v)
    for p in figuras:
        doc.add_heading(Path(p).stem.replace('_',' '),level=2); doc.add_picture(str(p),width=Inches(6.3))
    ruta=Path(ruta); ruta.parent.mkdir(parents=True,exist_ok=True); doc.save(ruta); return ruta

def generar_pdf_desde_docx(docx,pdf):
    import subprocess,shutil
    exe=shutil.which('libreoffice') or shutil.which('soffice')
    if not exe: raise RuntimeError('LibreOffice no está disponible para convertir a PDF.')
    pdf=Path(pdf); pdf.parent.mkdir(parents=True,exist_ok=True)
    subprocess.run([exe,'--headless','--convert-to','pdf','--outdir',str(pdf.parent),str(docx)],check=True,stdout=subprocess.PIPE,stderr=subprocess.PIPE)
    generado=pdf.parent/(Path(docx).stem+'.pdf')
    if generado!=pdf: generado.replace(pdf)
    return pdf
